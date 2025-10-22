from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
import yaml

def _mm(x): return Mm(float(x))

def _set_section_size(section, page_w_mm, page_h_mm, margins):
    section.page_width  = _mm(page_w_mm)
    section.page_height = _mm(page_h_mm)
    section.top_margin    = _mm(margins["top"])
    section.right_margin  = _mm(margins["right"])
    section.bottom_margin = _mm(margins["bottom"])
    section.left_margin   = _mm(margins["left"])

def _table_align(s: str):
    s = (s or "left").lower()
    return {"left": WD_TABLE_ALIGNMENT.LEFT,
            "center": WD_TABLE_ALIGNMENT.CENTER,
            "right": WD_TABLE_ALIGNMENT.RIGHT}.get(s, WD_TABLE_ALIGNMENT.LEFT)

def _para_align(s: str):
    s = (s or "left").lower()
    return {"left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT}.get(s, WD_ALIGN_PARAGRAPH.LEFT)

def _clear_table_borders(table):
    # Strip all table/cell borders via XML so there’s no rectangle
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.xpath("w:tblBorders")
    if borders:
        tblPr.remove(borders[0])
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def build_doc_flow(labels, config_path, out_docx):
    """
    Word: ONE STICKER PER PAGE (page size = label size).
    No borders. Times New Roman. Line sizes/weights from YAML.
    """
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    page   = cfg["page"]
    label  = cfg["label"]
    text   = cfg["text"]
    lines  = cfg.get("lines", [])
    word_c = cfg.get("word", {"align": "center"})

    page_w = float(label["width_mm"])
    page_h = float(label["height_mm"])
    margins = page.get("margin_mm", {"top": 2, "right": 2, "bottom": 2, "left": 2})

    doc = Document()
    _set_section_size(doc.sections[0], page_w, page_h, margins)

    table_alignment = _table_align(word_c.get("align", "center"))
    para_alignment  = _para_align(text.get("align", "center"))
    line_spacing    = text.get("line_spacing", 1.0)

    def add_line(paragraph, value, size_pt, bold=False):
        run = paragraph.add_run(value)
        run.font.name = text["font_name"]
        run.font.size = Pt(size_pt)
        run.bold = bool(bold)

    content_w_mm = page_w - float(margins["left"]) - float(margins["right"])
    content_h_mm = page_h - float(margins["top"]) - float(margins["bottom"])

    for i, lab in enumerate(labels):
        if i > 0:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            _set_section_size(sec, page_w, page_h, margins)

        table = doc.add_table(rows=1, cols=1)
        table.style = None                      # no default grid
        _clear_table_borders(table)             # force no borders
        table.alignment = table_alignment
        table.autofit = False
        try:
            table.columns[0].width = _mm(content_w_mm)
        except Exception:
            pass

        cell = table.cell(0, 0)
        cell.width = _mm(content_w_mm)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # exact height to fill content area
        try:
            from docx.enum.table import WD_ROW_HEIGHT_RULE
            row = table.rows[0]
            row.height = _mm(content_h_mm)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        except Exception:
            pass

        # line1 paragraph
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = line_spacing
        p.alignment = para_alignment

        # line1 (bold 14)
        l1cfg = lines[0] if len(lines) > 0 else {"show": True, "bold": True}
        if l1cfg.get("show", True):
            add_line(p, lab.get("line1", ""), text.get("line1_size_pt", 14), l1cfg.get("bold", True))

        # line2..line4
        spec = [("line2","line2_size_pt"), ("line3","line3_size_pt"), ("line4","line4_size_pt")]
        for pos, (fname, fsize_key) in enumerate(spec, start=2):
            cfg_line = lines[pos - 1] if len(lines) >= pos else {"show": True, "bold": (pos == 4)}
            if cfg_line.get("show", True):
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after  = Pt(0)
                p2.paragraph_format.line_spacing = line_spacing
                p2.alignment = para_alignment
                add_line(p2, lab.get(fname, ""), text.get(fsize_key, 12), cfg_line.get("bold", pos == 4))

    doc.save(out_docx)
